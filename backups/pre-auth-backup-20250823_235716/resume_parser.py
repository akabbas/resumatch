#!/usr/bin/env python3
"""
Resume Parser Module
Extracts text from PDF and Word documents, then uses AI to parse content into structured sections.
"""

import os
import re
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import tempfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not available. PDF parsing will not work.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word document parsing will not work.")

class ResumeParser:
    """Parse resumes from PDF and Word documents using AI text analysis"""
    
    def __init__(self):
        self.section_keywords = {
            'summary': ['summary', 'profile', 'objective', 'overview', 'introduction'],
            'experience': ['experience', 'work history', 'employment', 'professional background', 'career'],
            'education': ['education', 'academic', 'degree', 'university', 'college', 'school'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'technologies'],
            'certifications': ['certifications', 'certificates', 'accreditations', 'training'],
            'projects': ['projects', 'portfolio', 'achievements', 'key projects'],
            'languages': ['languages', 'language skills', 'fluency'],
            'interests': ['interests', 'hobbies', 'activities', 'volunteer']
        }
        
        # Common job titles and companies for extraction
        self.job_title_patterns = [
            r'(?i)(senior|junior|lead|principal|staff|associate|director|manager|engineer|analyst|developer|consultant|specialist|coordinator|administrator)',
            r'(?i)(software|data|business|systems|product|project|devops|cloud|security|network|database|web|full.?stack|front.?end|back.?end)',
            r'(?i)(engineer|developer|analyst|architect|manager|consultant|specialist|coordinator|administrator)'
        ]
        
        self.company_patterns = [
            r'(?i)(inc\.?|corp\.?|llc|ltd\.?|company|technologies|solutions|systems|group|partners)',
            r'(?i)(microsoft|google|amazon|apple|facebook|netflix|salesforce|oracle|ibm|intel)',
            r'(?i)(startup|enterprise|agency|consulting|consultancy)'
        ]
    
    def parse_resume(self, file_path: str) -> Dict:
        """
        Parse a resume file and extract structured information
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            Dict: Parsed resume data with sections
        """
        try:
            # Extract text from file
            text = self._extract_text(file_path)
            if not text:
                return {'error': 'Could not extract text from file'}
            
            # Parse text into sections
            parsed_data = self._parse_text_sections(text)
            
            # Extract additional metadata
            parsed_data.update(self._extract_metadata(text))
            
            logger.info(f"Successfully parsed resume: {len(parsed_data)} sections extracted")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            return {'error': f'Failed to parse resume: {str(e)}'}
    
    def _extract_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF, Word document, or text file"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx_text(file_path)
        elif file_ext == '.txt':
            return self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF using pdfplumber"""
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is required for PDF parsing")
        
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extract text from Word document using python-docx"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document parsing")
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + "\n"
            
            logger.info(f"Extracted {len(text)} characters from Word document")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting Word document text: {e}")
            return None
    
    def _extract_text_file(self, file_path: str) -> Optional[str]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"Extracted {len(text)} characters from text file")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text file: {e}")
            return None
    
    def _parse_text_sections(self, text: str) -> Dict:
        """Parse text into structured sections using AI-like analysis"""
        sections = {}
        
        # Split text into lines for analysis
        lines = text.split('\n')
        
        # Find section boundaries
        section_boundaries = self._find_section_boundaries(lines)
        
        # Extract content for each section
        for section_name, (start_idx, end_idx) in section_boundaries.items():
            if start_idx < end_idx:
                section_text = '\n'.join(lines[start_idx:end_idx]).strip()
                if section_text:
                    sections[section_name] = self._clean_section_text(section_text)
        
        return sections
    
    def _find_section_boundaries(self, lines: List[str]) -> Dict[str, Tuple[int, int]]:
        """Find the boundaries of different sections in the resume"""
        section_boundaries = {}
        current_section = None
        current_start = 0
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            detected_section = self._detect_section_header(line_lower)
            
            if detected_section:
                # Save previous section if exists
                if current_section:
                    section_boundaries[current_section] = (current_start, i)
                
                # Start new section
                current_section = detected_section
                current_start = i + 1
        
        # Save the last section
        if current_section:
            section_boundaries[current_section] = (current_start, len(lines))
        
        return section_boundaries
    
    def _detect_section_header(self, line: str) -> Optional[str]:
        """Detect if a line is a section header"""
        # Remove common formatting and punctuation
        clean_line = re.sub(r'[^\w\s]', '', line.lower())
        
        # Check against section keywords
        for section, keywords in self.section_keywords.items():
            for keyword in keywords:
                if keyword in clean_line:
                    return section
        
        # Check for common header patterns
        if re.match(r'^[A-Z\s]{3,}$', line):  # ALL CAPS lines
            return self._guess_section_from_caps(line)
        
        return None
    
    def _guess_section_from_caps(self, line: str) -> Optional[str]:
        """Guess section from ALL CAPS headers"""
        line_lower = line.lower()
        
        if any(word in line_lower for word in ['experience', 'work', 'employment']):
            return 'experience'
        elif any(word in line_lower for word in ['education', 'academic']):
            return 'education'
        elif any(word in line_lower for word in ['skills', 'expertise']):
            return 'skills'
        elif any(word in line_lower for word in ['summary', 'profile', 'objective']):
            return 'summary'
        elif any(word in line_lower for word in ['certifications', 'training']):
            return 'certifications'
        
        return None
    
    def _clean_section_text(self, text: str) -> str:
        """Clean and format section text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove bullet points and numbering
        text = re.sub(r'^[\s•\-\*\d\.]+', '', text, flags=re.MULTILINE)
        
        # Clean up line breaks
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text.strip()
    
    def _extract_metadata(self, text: str) -> Dict:
        """Extract metadata like job title, company, etc."""
        metadata = {}
        
        # Try to extract job title
        job_title = self._extract_job_title(text)
        if job_title:
            metadata['job_title'] = job_title
        
        # Try to extract company name
        company = self._extract_company_name(text)
        if company:
            metadata['company'] = company
        
        # Try to extract summary from first few lines
        summary = self._extract_summary(text)
        if summary:
            metadata['summary'] = summary
        
        # Try to extract skills
        skills = self._extract_skills(text)
        if skills:
            metadata['skills'] = skills
        
        return metadata
    
    def _extract_job_title(self, text: str) -> Optional[str]:
        """Extract job title from resume text"""
        lines = text.split('\n')
        
        # Look for job title in first few lines (usually near the top)
        for i, line in enumerate(lines[:10]):
            line_clean = re.sub(r'[^\w\s]', '', line.strip())
            
            # Check if line contains job title patterns
            for pattern in self.job_title_patterns:
                if re.search(pattern, line_clean, re.IGNORECASE):
                    # Clean up the extracted title
                    title = line.strip()
                    title = re.sub(r'^[^\w]*', '', title)  # Remove leading non-word chars
                    title = re.sub(r'[^\w\s]*$', '', title)  # Remove trailing non-word chars
                    
                    if len(title) > 3 and len(title) < 100:  # Reasonable length
                        return title
        
        return None
    
    def _extract_company_name(self, text: str) -> Optional[str]:
        """Extract company name from resume text"""
        lines = text.split('\n')
        
        # Look for company name near job title
        for i, line in enumerate(lines[:15]):
            line_clean = re.sub(r'[^\w\s]', '', line.strip())
            
            # Check if line contains company patterns
            for pattern in self.company_patterns:
                if re.search(pattern, line_clean, re.IGNORECASE):
                    # Clean up the extracted company name
                    company = line.strip()
                    company = re.sub(r'^[^\w]*', '', company)
                    company = re.sub(r'[^\w\s]*$', '', company)
                    
                    if len(company) > 2 and len(company) < 80:  # Reasonable length
                        return company
        
        return None
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract summary/profile section from resume"""
        lines = text.split('\n')
        
        # Look for summary in first few lines
        summary_lines = []
        in_summary = False
        
        for i, line in enumerate(lines[:20]):
            line_lower = line.lower().strip()
            
            # Check if we're entering summary section
            if any(keyword in line_lower for keyword in ['summary', 'profile', 'objective']):
                in_summary = True
                continue
            
            # Check if we're leaving summary section
            if in_summary and any(keyword in line_lower for keyword in ['experience', 'education', 'skills']):
                break
            
            if in_summary and line.strip():
                summary_lines.append(line.strip())
        
        if summary_lines:
            summary = ' '.join(summary_lines)
            # Limit summary length
            if len(summary) > 500:
                summary = summary[:500] + "..."
            return summary
        
        return None
    
    def _extract_skills(self, text: str) -> Optional[str]:
        """Extract skills from resume text"""
        # Look for skills section
        skills_match = re.search(r'(?i)skills?[:\s]*(.*?)(?=\n\s*[A-Z]|\n\s*\n|$)', text, re.DOTALL)
        
        if skills_match:
            skills_text = skills_match.group(1).strip()
            
            # Clean up skills text
            skills_text = re.sub(r'[•\-\*]', ',', skills_text)  # Replace bullets with commas
            skills_text = re.sub(r'\s+', ' ', skills_text)  # Normalize whitespace
            skills_text = re.sub(r',\s*,', ',', skills_text)  # Remove empty commas
            skills_text = skills_text.strip(', ')  # Remove leading/trailing commas
            
            if skills_text and len(skills_text) > 5:
                return skills_text
        
        return None

def parse_resume_file(file_path: str) -> Dict:
    """Convenience function to parse a resume file"""
    parser = ResumeParser()
    return parser.parse_resume(file_path)

if __name__ == "__main__":
    # Test the resume parser
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python resume_parser.py <resume_file>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)
    
    # Parse the resume
    result = parse_resume_file(file_path)
    
    if 'error' in result:
        print(f"Error: {result['error']}")
        sys.exit(1)
    
    # Print results
    print("Resume parsed successfully!")
    print("=" * 50)
    
    for section, content in result.items():
        if content:
            print(f"\n{section.upper()}:")
            print("-" * 30)
            if isinstance(content, str) and len(content) > 100:
                print(content[:100] + "...")
            else:
                print(content)
